import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx not found.")
    sys.exit(1)

def add_markdown_formatted_text(paragraph, text):
    """
    Parses simple markdown (bold **, italic *, superscript [^n]) and adds runs to paragraph.
    Note: Does not handle nested formatting perfectly, but sufficient for this mockup.
    """
    # Regex to split format tokens:
    # Captures: **bold**, *italic*, [^ref]
    # We use a combined pattern. 
    pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|\[\^.*?\])'
    tokens = re.split(pattern, text)
    
    for token in tokens:
        if not token: continue
        
        bold = False
        italic = False
        superscript = False
        content = token
        
        if token.startswith('***') and token.endswith('***'):
            bold = True
            italic = True
            content = token[3:-3]
        elif token.startswith('**') and token.endswith('**'):
            bold = True
            content = token[2:-2]
        elif token.startswith('*') and token.endswith('*'):
            italic = True
            content = token[1:-1]
        elif token.startswith('[^') and token.endswith(']'):
            superscript = True
            content = token[2:-1]
        
        run = paragraph.add_run(content)
        run.font.name = 'Times New Roman'
        run.bold = bold
        run.italic = italic
        run.font.superscript = superscript

def create_docx(md_path, docx_path):
    document = Document()
    
    # Styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    
    with open(md_path, 'r') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line: continue
        
        # Headers (Assuming # is H1, ## is H2, etc)
        if line.startswith('# '):
            p = document.add_heading(level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0,0,0)
            run.bold = False 
            continue
            
        if line.startswith('## '):
            p = document.add_heading(level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[3:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0,0,0)
            run.bold = False
            continue
            
        if line.startswith('### '):
            p = document.add_heading(level=3)
            # SQ often uses caps or italics. 
            run = p.add_run(line[4:])
            run.font.name = 'Times New Roman'
            run.font.italic = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0,0,0)
            continue
            
        # Separator
        if line == "***":
             p = document.add_paragraph()
             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
             p.add_run("***")
             continue

        # Blockquotes/Figures (> ...)
        if line.startswith('>'):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            # Remove >
            content = line[1:].strip()
            
            # Simple heuristic: if it looks like a figure caption, maybe bold the label?
            # Our helper handles **Figure 1** if marked in MD.
            add_markdown_formatted_text(p, content)
            continue

        # Standard Paragraph
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.5)
        
        # Special handling for Metadata lines (Author, Abstract)
        # If line is bold and centered-like (short)
        if line.startswith('**') and line.endswith('**') and len(line) < 60:
             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
             p.paragraph_format.first_line_indent = Inches(0)
        
        add_markdown_formatted_text(p, line)
        
    document.save(docx_path)
    print(f"Generated {docx_path}")

if __name__ == "__main__":
    create_docx(
        "/Users/joelfarthing/Library/Mobile Documents/com~apple~CloudDocs/Documents/Filament_Local_Mac/shakespeare_quarterly_mockup.md",
        "/Users/joelfarthing/Library/Mobile Documents/com~apple~CloudDocs/Documents/Filament_Local_Mac/shakespeare_quarterly_mockup.docx"
    )
